import io
import re
from typing import List, Dict, Tuple
from xml.sax.saxutils import escape
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from ..models.document import GeneratedDocument

class RegulatoryPDFExporter:
    def __init__(self, session_id: str = None):
        self.session_id = session_id
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom styles for regulatory documents"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='RegulatoryTitle',
            parent=self.styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=HexColor('#1f4e79')
        ))
        
        # Section header style
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            textColor=HexColor('#1f4e79')
        ))
        
        # Regular content style
        self.styles.add(ParagraphStyle(
            name='RegulatoryContent',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            firstLineIndent=20
        ))

    def generate_regulatory_document(self, document: GeneratedDocument) -> bytes:
        """Generate a PDF document from a GeneratedDocument model"""
        content = document.dict()  # Convert Pydantic model to dict for processing
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # Document title
        story.append(Paragraph(content.get('title', 'Regulatory Document'), self.styles['RegulatoryTitle']))
        story.append(Spacer(1, 0.5*inch))
        
        # Add generation metadata
        story.append(Paragraph(f"Generated on: {content.get('generated_at', '')}", self.styles['Normal']))
        story.append(Paragraph(f"Session ID: {content.get('session_id', '')}", self.styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Add sections
        for section in content.get('sections', []):
            # Section title
            story.append(Paragraph(section.get('title', ''), self.styles['SectionHeader']))
            
            # Section content
            content_text = section.get('content', '')
            # Clean and escape the content
            content_text = self._clean_content(content_text)
            story.append(Paragraph(content_text, self.styles['RegulatoryContent']))
            
            # Add source count info
            source_count = section.get('source_count', 0)
            if source_count > 0:
                story.append(Paragraph(f"<i>Sources referenced: {source_count}</i>", self.styles['Normal']))
            
            story.append(Spacer(1, 0.2*inch))
        
        # Build the PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _clean_content(self, content: str) -> str:
        """Clean and format content for PDF generation"""
        if not content:
            return ""
        
        # Escape HTML entities
        content = escape(content)
        
        # Replace line breaks with paragraph breaks
        content = content.replace('\n\n', '<br/><br/>')
        content = content.replace('\n', '<br/>')
        
        return content

class RegulatoryDOCXExporter:
    def __init__(self, session_id: str = None):
        self.session_id = session_id

    def generate_regulatory_document(self, document: GeneratedDocument) -> bytes:
        """Generate a DOCX document from a GeneratedDocument model"""
        content = document.dict()  # Convert Pydantic model to dict for processing
        
        # Create a new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add custom styles
        self._setup_custom_styles(doc)
        
        # Document title
        title = doc.add_heading(content.get('title', 'Regulatory Document'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Generated on: {content.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}\n")
        metadata_para.add_run(f"Session ID: {content.get('session_id', self.session_id or 'N/A')}")
        
        # Add a page break
        doc.add_page_break()
        
        # Add sections
        for section in content.get('sections', []):
            # Section title
            section_title = doc.add_heading(section.get('title', ''), level=1)
            
            # Section content
            content_text = section.get('content', '')
            content_text = self._clean_content_for_docx(content_text)
            
            # Split content by paragraphs and add them
            paragraphs = content_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.style = 'Normal'
            
            # Add source count info
            source_count = section.get('source_count', 0)
            if source_count > 0:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(f"Sources referenced: {source_count}")
                source_run.italic = True
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_custom_styles(self, doc):
        """Setup custom styles for regulatory documents"""
        styles = doc.styles
        
        # Create a custom normal style if it doesn't exist
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
        except KeyError:
            pass
        
        # Create custom heading styles
        try:
            heading1_style = styles['Heading 1']
            heading1_style.font.name = 'Calibri'
            heading1_style.font.size = Pt(14)
            heading1_style.font.bold = True
        except KeyError:
            pass

    def _clean_content_for_docx(self, content: str) -> str:
        """Clean and format content for DOCX generation"""
        if not content:
            return ""
        
        # Remove HTML tags if any
        content = re.sub(r'<[^>]+>', '', content)
        
        # Clean up extra whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = content.strip()
        
        return content

class DocumentExporter:
    """Unified document exporter that handles both PDF and DOCX formats"""
    
    def __init__(self, session_id: str = None):
        self.session_id = session_id
        self.pdf_exporter = RegulatoryPDFExporter(session_id)
        self.docx_exporter = RegulatoryDOCXExporter(session_id)
    
    def export_document(self, document: GeneratedDocument, format_type: str = 'pdf') -> bytes:
        """Export document in specified format"""
        if format_type.lower() == 'pdf':
            return self.pdf_exporter.generate_regulatory_document(document)
        elif format_type.lower() == 'docx':
            return self.docx_exporter.generate_regulatory_document(document)
        else:
            raise ValueError(f"Unsupported format: {format_type}. Supported formats: pdf, docx")
